import io
import re

from docx import Document
from pypdf import PdfReader

MAX_FILE_BYTES = 5 * 1024 * 1024
MIN_TEXT_CHARS = 200
MAX_TEXT_CHARS = 30_000


class ExtractionError(Exception):
    """Raised when a CV file cannot be turned into usable text."""


def extract_text(filename: str, data: bytes) -> str:
    if len(data) > MAX_FILE_BYTES:
        raise ExtractionError("File is larger than 5 MB — please upload a smaller CV.")

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        text = _from_pdf(data)
    elif ext == "docx":
        text = _from_docx(data)
    else:
        raise ExtractionError("Unsupported file type — please upload a PDF or DOCX.")

    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if len(text) < MIN_TEXT_CHARS:
        raise ExtractionError(
            "Could not extract readable text from this file. "
            "If it is a scanned PDF, please upload a text-based version."
        )
    return text[:MAX_TEXT_CHARS]


def _from_pdf(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError("Could not read this PDF file.") from exc


def _from_docx(data: bytes) -> str:
    try:
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)
    except Exception as exc:
        raise ExtractionError("Could not read this DOCX file.") from exc
